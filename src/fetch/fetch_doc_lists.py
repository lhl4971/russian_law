import os
import re
import json
import requests
from docx import Document

def parse_file_list_docx(filepath):
    doc = Document(filepath)

    # 读取标题（页眉）
    header = doc.sections[0].header
    title = "\n".join(p.text.strip() for p in header.paragraphs if p.text.strip())

    # 读取正文段落，去掉空白
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 拼接输出
    text = []

    # 主体编号部分
    for para in paragraphs:
        text.append(para)

    return f"{title}\n\n" + "\n".join(text)


def fetch_doc_lists(input_data, state_duty, review_period, collection, application_type=None):        
    results = {}
    headers = {"User-Agent": "Mozilla/5.0"}
    os.makedirs(f"data/raw/{collection}", exist_ok=True)

    # 定义模式到 application_type 的映射
    patterns = [
        (r"\bВНЖ\b", "ВНЖ"),
        (r"\bРВПО\b", "РВПО"),
        (r"\bРВП\b", "РВП")
    ]

    for item in input_data:
        url = item["href"]
        filename = url.split("/")[-1]
        filepath = os.path.join(f"data/raw/{collection}", filename)

        # 下载文件（若不存在）
        if not os.path.exists(filepath):
            r = requests.get(url, headers=headers, verify=False)
            r.raise_for_status()
            with open(filepath, "wb") as f:
                f.write(r.content)
            print(f"下载完成: {filename}")

        required_documents_list = parse_file_list_docx(filepath)

        # 匹配 application_type
        for pattern, app_type in patterns:
            if re.search(pattern, item["text"]):
                application_type = app_type
                break

        if not application_type:
            raise ValueError(f"无法识别申请类型: {item['text']}")

        # 初始化 application_type
        if application_type not in results:
            results[application_type] = []

        # 添加结果
        results[application_type].append({
            "id": len(results[application_type]),
            "application_type": application_type,
            "text": item["text"],
            "required_documents_list": required_documents_list,
            "state_duty_law": state_duty[application_type]["law"],
            "receipt_form_payment": state_duty[application_type]["href"],
            "review_period": review_period[application_type]["review_period"]
        })

    return results

if __name__ == "__main__":
    parsed_doc_lists = {}
    with open("data/processed/list_and_blanks/state_duty.json", "r", encoding="utf-8") as f:
        state_duty = json.load(f)
    with open("data/processed/list_and_blanks/review_period.json", "r", encoding="utf-8") as f:
        review_period = json.load(f)

    # РВП и ВЖ
    with open("data/processed/list_and_blanks/list.json", "r", encoding="utf-8") as f:
        data = json.load(f)
    results = fetch_doc_lists(data, state_duty, review_period, "trp_rp")
    parsed_doc_lists.update(results)

    # Гражданство
    with open("data/processed/list_and_blanks/list_citizenship.json", "r", encoding="utf-8") as f:
        data = json.load(f)
    results = fetch_doc_lists(data, state_duty, review_period, "citizenship", application_type="Гражданство")
    parsed_doc_lists.update(results)

    # 保存结果
    with open("data/processed/list_and_blanks/parsed_doc_lists.json", "w", encoding="utf-8") as f:
        json.dump(parsed_doc_lists, f, ensure_ascii=False, indent=2)

    print("解析完成，结果已保存到 data/processed/list_and_blanks/parsed_doc_lists.json")
